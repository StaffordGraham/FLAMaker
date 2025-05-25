import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
from tkcalendar import Calendar
from Person import Person
from CaseDetails import CaseDetails
from datetime import datetime
from tkcalendar import DateEntry
import os
import subprocess
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Cm
from docx.shared import Inches
import functions
from functions import article
from datetime import date, timedelta

import sys




def write_order(case_details):
        #if case_details.definitions:
            #family_home_def=case_details.definitions[0]
        court_name=f"In the Family Court\n sitting at {case_details.court_name}"
        case_number_text=f"Case No: {case_details.case_number}"
        appapp=case_details.application
        appo =f"{appapp} \n Family Law Act 1996"
            
      
         #CREATE DOCUMENT   
        doc=Document()
        #COURT AND CASE NUMBER
        heading_table=doc.add_table(1,3)
        coat_of_arms_cell =heading_table.cell(0,0)
        court_name_cell=heading_table.cell(0,1)
        case_number_cell=heading_table.cell(0,2)
        coat_of_arms_cell.paragraphs[0].add_run().add_picture('Coat_of_arms.png',width=Inches(1.25))
        court_name_paragraph=court_name_cell.paragraphs[0]
        court_name_run=court_name_paragraph.add_run(court_name)
        court_name_run.bold=True
        
        case_number_paragraph=case_number_cell.paragraphs[0]
        case_number_run=case_number_paragraph.add_run(case_details.case_number)
        case_number_run.bold=True
        
        #LINE
        
        line_paragraph = doc.add_paragraph()
        line_paragraph.add_run('_'*100)
        #THE APPLICATION AND THE STATUTE
        paragraph=doc.add_paragraph()
        appo=case_details.application
        pprun=paragraph.add_run(appo)
        pprun.bold=True
        pparagraph=doc.add_paragraph()
        pprun=pparagraph.add_run("Family Law Act 1996")
        pprun.bold=True
        paragraph=doc.add_paragraph()
        
        #LINE
        line_paragrap=doc.add_paragraph()  # Adjust the number for length
        line_paragrap.add_run('_'*100)
        #TABLE FOR CHILDREN
        kidno =len(case_details.children)
        
        if kidno>0:
            
            if kidno>1:
                label_children="The full names of the children"
                label_date="Dates of Birth"
            else:
                label_children="The full name of the child"
                label_date="Date of Birth"
        child_table=doc.add_table(rows=kidno+1,cols=3)
        child_table.columns[0].width=Cm(20)
        child_table.columns[1].width=Cm(3)
        child_table.columns[2].width=Cm(3)
       
        cell_00=child_table.cell(0,0)
        cell_00.text=label_children
        cell_01=child_table.cell(0,1)
        cell_01.text="Boy or Girl"
        cell_02=child_table.cell(0,2)
        cell_02.text=label_date
        
        #ADD CHILDREN INFO
        
        for i, child in enumerate(case_details.children, start=1):  # start=1 to fill rows below the header
            child_full_name =child.first_name + " "+child.last_name
            child_table.cell(i, 0).text = child_full_name  # Name
            child_table.cell(i, 1).text = child.gender # Gender
            child_table.cell(i, 2).text = child.full_date_of_birth  # Date of Birth (convert to string if necessary)
        
        doc.add_paragraph()
        
        #THE WARNING
        
        warning_text=[] 
        warning_text.append(f"IMPORTANT NOTICE TO THE RESPONDENT {case_details.respondent.first_name.upper()} ")
        warning_text.append(case_details.respondent.last_name.upper())
        warning_text.append(f" of {case_details.respondent.address_building_and_street.upper()}, ")
        warning_text.append(f"{case_details.respondent.address_second_line.upper()} ")
        warning_text.append(f"{case_details.respondent.address_town_or_city.upper()} ")
        warning_text.append(f"{case_details.respondent.address_postcode.upper()}.\n")
        warning_text.append("YOU MUST OBEY THIS ORDER. You should read it carefully .If you do not understand anything in this order you should go to a solicitor, Legal Advice Centre or Citizens Advice Bureau. You have the right to apply to the court to change or cancel the order.\n")
        warning_text.append("WARNING: IF YOU DISOBEY THIS ORDER, YOU MAY BE HELD TO BE IN CONTEMPT OF COURT AND MAY BE IMPRISONED, FINED, OR HAVE YOUR ASSETS SEIZED")
        final_warning_text=' '.join(warning_text) 
        
        
        warning_table=doc.add_table(rows=1,cols=1)
        warning_table.autofit=True
        warning_table.style='Table Grid'
        
        warning_cell =warning_table.cell(0,0)
        warning_cell_paragraph=warning_cell.paragraphs[0]
        run = warning_cell_paragraph.add_run(final_warning_text)
        run.bold=True
        doc.add_paragraph()
        doc.add_paragraph()
        
        #THE JUDGE, HEARING DATE & TYPE
        judge_date_hearing_type=[]
        judge_date_hearing_type.append('Before')
        if case_details.judge_rank:
            judge_date_hearing_type.append(case_details.judge_rank)
            
        if case_details.judge_name:
            judge_date_hearing_type.append(case_details.judge_name)
          
        if case_details.hearing_date_object:
            judge_date_hearing_type.append(' in private on ')
            the_hearing_date=case_details.hearing_date_object.strftime("%d %B %Y")      
            judge_date_hearing_type.append(the_hearing_date)
            
        if case_details.listed_for:
            judge_date_hearing_type.append(' at')
            lfor =case_details.listed_for
            lfor=lfor.lower()
            artic= article(case_details.listed_for)
            judge_date_hearing_type.append(artic+' '+lfor)
            
       
            
            
        
        judge_date_hearing_type.append('.')
        judge_and_hearing=" ".join(judge_date_hearing_type)
        doc.add_paragraph(judge_and_hearing)        
        #THE PARTIES
        
        
        parties_definitions_text=[]
        parties_definitions_text.append(f"The applicant is {case_details.applicant.first_name} {case_details.applicant.last_name}")
        reppo=case_details.app_rep
        beppo=case_details.resp_rep
        
        parties_definitions_text.append(f"represented by {case_details.app_rep.full_name}")
        parties_definitions_text.append(f"The respondent is {case_details.respondent.first_name} {case_details.respondent.last_name}")
        parties_definitions_text.append(f"represented by {case_details.resp_rep.full_name}")
        the_parties_defs=' '.join(parties_definitions_text)
        parties_table=doc.add_table(rows=2,cols=2)
        for row in parties_table.rows:
            row.cells[0].width = Cm(3.0)  # Adjust as needed
            row.cells[1].width = Cm(14.0)  # 
    
        parties_label=parties_table.cell(0,0)
        parties_para =parties_label.add_paragraph()
        run =parties_para.add_run("The Parties: ")
        run.bold=True
        parties_cell=parties_table.cell(0,1)
        parties_cell_para=parties_cell.add_paragraph()
        run=parties_cell_para.add_run(the_parties_defs)
        parties_cell_para.alignment =0 # Left-aligned
               
         #DEFINITIONS - CHILDREN
        definitions_table=doc.add_table(rows=2,cols=2)
        definitions_label=definitions_table.cell(0,0)
        definit_head=definitions_label.add_paragraph()
        run=definit_head.add_run("Definitions")
        run.bold=True
        
        numb_children = len(case_details.children)
        
        
        
        if numb_children>1:
            chilo="The “relevant children” within the meaning of Family Law Act 1996 are:"
            paragraph_text=chilo
            
            starting_char='a'
            
            for index, child in enumerate(case_details.children):
                    sentence=child.sentence
                    prefix =chr(ord(starting_char)+index)
                    paragraph_text+=f"\n  {prefix}\t{sentence}\n"
                    
            doc.add_paragraph(paragraph_text,style='ListNumber')
            
        else:
            chilo=f"The relevant child within the meaning of the Family Law Act 1995 is {case_details.children[0].sentence}"
            doc.add_paragraph(chilo,style='ListNumber')

        
        
            
            #DEFINITIONS FAMILY HOME
            
        fam_home_stringbuilder =[]
        fam_home_stringbuilder.append('The "family home" is the property at')
        fam_home_stringbuilder.append(case_details.applicant.address_building_and_street)
        fam_home_stringbuilder.append(case_details.applicant.address_second_line)
        fam_home_stringbuilder.append(case_details.applicant.address_town_or_city)
        fam_home_stringbuilder.append(case_details.applicant.address_postcode)
        family_home_def=' '.join(fam_home_stringbuilder)
        doc.add_paragraph(family_home_def  ,style='ListNumber')
            
        #RECITALS

        if len(case_details.recitals) >0:
            recitals_para_head=doc.add_paragraph()
            rec_run=recitals_para_head.add_run("Recitals")
            rec_run.bold=True
        
            for recit_para in case_details.recitals:
                doc.add_paragraph(recit_para,style='ListNumber')
               
        
        #UNDERTAKINGS
        if len(case_details.undertakings)>0:
            undertakings_para_head=doc.add_paragraph()
            ut_rn=undertakings_para_head.add_run('Undertakings')
            ut_rn.bold=True
            for upara in case_details.undertakings:
                doc.add_paragraph(upara,style='ListNumber')
            
        
        
         #ORDERS
            
        if case_details.consent_order ==True:
            orders_para_head="IT IS ORDERED BY CONSENT"
        else:
            orders_para_head="IT IS ORDERED"
        
        app_name=[]
        app_name1 =case_details.applicant.first_name
        app_name2=case_details.applicant.last_name
        app_name=app_name1+" "+app_name2
        resp_name1=case_details.respondent.first_name
        resp_name2=case_details.respondent.last_name
        resp_name=resp_name1+" "+resp_name2
              
        mat_home1=case_details.applicant.address_building_and_street
        mat_home2=case_details.applicant.address_second_line
        mat_home3=case_details.applicant.address_town_or_city
        mat_home4=case_details.applicant.address_postcode
        mat_home=mat_home1+", "+mat_home2+", "+mat_home3+", "+mat_home4
        relat=case_details.relationship
        if len(case_details.children)==1:
            kidno="child"
            
        if len (case_details.children)>1:
            kidno="children"
            
        school_address= ""
            
        
        tody=date.today()
        leave_date = tody + timedelta(days=7)
        
        ord_dict={}
        key = "Use or threaten violence"
        value =(
            f"The respondent {resp_name} must not use or threaten violence against the applicant {app_name}"
         f" and must not instruct, encourage or in any way suggest any other person should do so."
         )
        
        ord_dict[key]=value
        
        key="Not intimidate,harass or pester"
        value=(f"The respondent {resp_name} must not intimidate, harass or pester the applicant {app_name} "
                f" and must not instruct, encourage or in any way suggest any other person should do so."
        )
        ord_dict[key]=value
        
        key="Not communicate"

        value =(f"The respondent, {resp_name}, must not telephone, text, email or otherwise contact"
                f" the applicant {app_name}, including via social networking websites or other forms of "
                f"electronic messaging."   
        )
        ord_dict[key]=value
        
        key="Not damage property"
        value=(f"The respondent, {resp_name}, must not damage, attempt to damage or threaten to damage"
               f"any property owned by or in the possession or controlof the applicant {app_name},"
               f"and must not instruct,encourage or in any way suggest that any other person should do so."
        )
        ord_dict[key]=value
        
        key="Not damage family home"
        value =(f"The respondent {resp_name}, must not damage, attempt to damage or threaten to damage"
                f"the property or contents of {mat_home} and must not instruct, encourage or in any way suggest"
                f"that any other person should do so."
              
        )
        ord_dict[key]=value
        
        key ="Not Enter family home"

        value=(f"The respondent {resp_name}, must not go to, enter or attempt to enter {mat_home} or "
               f" any property where {case_details.respondent.nominative} believes the applicant {app_name}"
               f" to be living."
        )
        
        ord_dict[key]=value
        key="Non mol zonal"
        ord_dict[key]=value
        value=(f" The respondent {resp_name} must not approach within 100 meters of the {mat_home}"
        )
        
        key="Not intimidate, harass or pester children"

        value =(f"The Respondent {resp_name}, must not use or threaten violence against the relevant {kidno} "
                f" and must not instruct, encourage or in any way suggest that any other person should do so."
        )
        ord_dict[key]=value
        
        key="Not communicate with children"

        value=(f"The respondent {resp_name}, must not telephone, text,email or otherwise contact or attempt  "
               f"the relevant {kidno} including via social networking websites or other forms of electronic"
               f" messaging."
        )
        ord_dict[key]=value
        
        key="Not enter children's school"
        value=(f"The respondent {resp_name}, must not go to, enter or attempt to enter the school premises"
            f" known as {school_address} except by prior written invitation from the school authorities."
        )
        ord_dict[key]=value
        
        
        key = "Declaration of Mat Home Rights"
        value =(
            f"The court declares that the applicant "
         f"{app_name}, has home rights in {mat_home}."
         )
        
        ord_dict[key]=value
        
        key="Mat Home Rights Survive Divorce"
        value=(f"The court declares that the {app_name}'s "
        f"home rights in {mat_home} shall not end when the respondent {resp_name}"
        f" dies or their {relat.lower()} is dissolved and shall continue until "
        f"the determination of the applicant's financial provision claims or a further order is made."
        )
        ord_dict[key]=value
        
        key="R Allow Occupy  of Mat Home"
        value =(f"The respondent, {resp_name}, shall allow the applicant, {app_name}, "
        f"to occupy the property at {mat_home}."
    
        )
        ord_dict[key]=value
        
        key="R Must Not Occupy Mat Home"
        value=(f"The respondent, {resp_name}, must not occupy the property at {mat_home}"
        )
        ord_dict[key]=value
        
        key="R leave Mat Home"
        value =(f"The respondent {resp_name}, shall leave the property at {mat_home}  by "
        f"4:00 pm on {leave_date}"
        )
        ord_dict[key]=value
        
        key ="R Not Return to Mat Home"
        value=(f"Having left {mat_home}, the respondent {resp_name}, "
        "must not return to, enter or attempt to enter it."
        )
        ord_dict[key]=value
        
        key="Not Obstruct A's Occupation"
        value =(f"The Respondent {resp_name}, must not obstruct, harass, or interfere with the "
        f"applicant {app_name}'s peaceful occupation of {mat_home}"
        )
        ord_dict[key]=value
        
        key="Right of Entry"
        value=(f"The applicant {app_name} has the right to enter into and occupy  the {mat_home} "
        f" and the respondent {resp_name},shall allow the applicant to do so."
        )
        ord_dict[key]=value
        
        key="Right Not to be Evicted"
        value=(f"The applicant {app_name}, has the right not to be evicted or excluded from, and the "
        f"respondent {resp_name}, must not evict or exclude the applicant from {mat_home}"
        )
        ord_dict[key]=value
        
        for key in case_details.order_choices:
            case_details.orders.append(ord_dict[key])
                
        oph=doc.add_paragraph()
        ord_head_run=oph.add_run(orders_para_head)
        ord_head_run.bold=True
        for opara in case_details.orders:
            doc.add_paragraph(opara,style='ListNumber')
            
        doc.add_paragraph("_" * 5)
        judge=case_details.judge_rank+" "+case_details.judge_name
            
            
        end=doc.add_paragraph()
            
        end_run =end.add_run(judge)
        doc.add_paragraph()
        date_end=doc.add_paragraph()
        the_end_date=case_details.hearing_date_object.strftime("%d %B %Y")
        run_date_end=date_end.add_run(the_end_date)


              
        # Save the document
        doc.save("family_court_document.docx")
        
        
        
        file_path=f"{case_details.file_name}.docx"
        doc.save(file_path)
        os.name='nt'
        os.startfile(file_path)

    